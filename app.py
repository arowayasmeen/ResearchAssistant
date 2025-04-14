from flask import Flask, render_template, make_response, jsonify, request, redirect, url_for, flash, session, Response
import time
from docx import Document
import io
import pdfkit  # Simpler PDF generation

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'  # Used for sessions and flash messages

# Mock user database for demonstration (in production, use a real database)
users = {}

@app.route('/')
def index():
    return render_template('index.html', page_title='Home')

@app.route('/literature')
def literature():
    return render_template('literature.html', page_title='Literature')

@app.route('/ideas')
def ideas():
    return render_template('ideas.html', page_title='Ideas')

@app.route('/paper')
def paper():
    return render_template('paper.html', page_title='Paper')

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        name = request.form['name']
        # Simple validation
        if email in users:
            flash('Email already exists', 'error')
            return redirect(url_for('signup'))
        # Store user (in a real app, you'd hash the password and use a database)
        users[email] = {
            'password': password,
            'name': name
        }
        flash('Account created successfully! Please log in.', 'success')
        return redirect(url_for('login'))
    return render_template('signup.html', page_title='Sign Up')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        # Validate credentials
        if email in users and users[email]['password'] == password:
            # Set session
            session['user_email'] = email
            session['user_name'] = users[email]['name']
            flash('Logged in successfully!', 'success')
            return redirect(url_for('index'))
        else:
            flash('Invalid email or password', 'error')
            return render_template('login.html', page_title='Log In')

@app.route('/profile')
def profile():
    # Check if user is logged in
    if 'user_email' not in session:
        flash('Please log in to view your profile', 'error')
        return redirect(url_for('login'))
    return render_template('profile.html', page_title='Profile')

@app.route('/logout')
def logout():
    # Clear the session
    session.pop('user_email', None)
    session.pop('user_name', None)
    flash('Logged out successfully', 'success')
    return redirect(url_for('index'))

# API endpoints for literature search
@app.route('/api/search', methods=['POST'])
def search_literature():
    """
    Endpoint to search for literature
    This is a placeholder that you would replace with actual AI integration
    """
    # Get search parameters from request
    data = request.json
    topic = data.get('topic', '')
    model = data.get('model', '')
    engine = data.get('engine', '')
    year_from = data.get('year_from')
    year_to = data.get('year_to')
    # Simulate processing time
    time.sleep(1)
    # In a real implementation, you would:
    # 1. Call your AI service or search API
    # 2. Process the results
    # 3. Return formatted data
    # Placeholder response
    return jsonify({
        'status': 'success',
        'message': 'This endpoint will be connected to your AI database',
        'query': {
            'topic': topic,
            'model': model,
            'engine': engine,
            'year_from': year_from,
            'year_to': year_to
        },
        'results': []  # This would contain actual results
    })

# API endpoint for idea generation
@app.route('/api/generate-ideas', methods=['POST'])
def generate_ideas():
    """
    Endpoint to generate research ideas
    This is a placeholder that you would replace with actual AI integration
    """
    # Get parameters from request
    data = request.json
    topic = data.get('topic', '')
    model = data.get('model', '')
    count = data.get('count', 5)
    idea_type = data.get('idea_type', '')
    # Simulate processing time
    time.sleep(1)
    # In a real implementation, you would:
    # 1. Call your AI service
    # 2. Process the generated ideas
    # 3. Return formatted data
    # Placeholder response
    return jsonify({
        'status': 'success',
        'message': 'This endpoint will be connected to your AI idea generation system',
        'query': {
            'topic': topic,
            'model': model,
            'count': count,
            'idea_type': idea_type
        },
        'ideas': []  # This would contain actual generated ideas
    })

# Helper function to clean HTML content
def clean_html_content(html_content):
    """Simple function to clean HTML content for export"""
    # In a real app, you might want to use a library like BeautifulSoup
    # to properly parse and clean the HTML
    # This is a simple placeholder
    return html_content

@app.route('/export_doc', methods=['POST'])
def export_doc():
    try:
        # Get the HTML content from the request
        html_content = request.form.get('document_content', '')
        
        # Create a new Document object for Word export
        doc = Document()
        
        # Use BeautifulSoup for better HTML parsing
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract and add title (h1)
        title_tag = soup.find('h1')
        if title_tag:
            doc.add_heading(title_tag.get_text(), 0)
            title_tag.extract()  # Remove title after processing
        
        # Extract author if present
        author_tag = soup.find('p', class_='author')
        if author_tag:
            doc.add_paragraph(author_tag.get_text()).italic = True
            author_tag.extract()  # Remove author after processing
        
        # Process the rest of the document
        for element in soup.find_all(['h2', 'p', 'ol', 'ul', 'li']):
            if element.name == 'h2':
                # Add section headings
                doc.add_heading(element.get_text(), 1)
            elif element.name == 'p':
                # Add paragraphs, checking for special formatting
                if 'section-note' in element.get('class', []):
                    paragraph = doc.add_paragraph(element.get_text())
                    paragraph.style = 'Intense Quote'  # Apply a style to section notes
                else:
                    doc.add_paragraph(element.get_text())
            elif element.name == 'ol':
                # Skip the ol tag itself and process its list items separately
                continue
            elif element.name == 'ul':
                # Skip the ul tag itself and process its list items separately
                continue
            elif element.name == 'li' and element.parent and element.parent.name == 'ol':
                # Add numbered list items with proper indentation
                doc.add_paragraph(element.get_text(), style='List Number')
            elif element.name == 'li' and element.parent and element.parent.name == 'ul':
                # Add bulleted list items with proper indentation
                doc.add_paragraph(element.get_text(), style='List Bullet')
        
        # Save the document
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        response = make_response(output.read())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = 'attachment; filename=research_paper.docx'
        
        return response
    except Exception as e:
        print(f"Error exporting document: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': f'Failed to export document: {str(e)}'
        }), 500

@app.route('/export_pdf', methods=['POST'])
def export_pdf():
    try:
        # Get the HTML content from the request
        html_content = request.form.get('document_content', '')
        
        # Clean the HTML content
        cleaned_content = clean_html_content(html_content)
        
        # Add basic CSS and wrap in proper HTML
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Research Paper</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 1in; }}
                h1 {{ font-size: 18pt; text-align: center; }}
                h2 {{ font-size: 14pt; }}
                p {{ font-size: 12pt; line-height: 1.5; }}
            </style>
        </head>
        <body>
            {cleaned_content}
        </body>
        </html>
        """
        
        # Create PDF using pdfkit - much simpler than WeasyPrint
        pdf = pdfkit.from_string(html, False)
        
        # Create response
        response = make_response(pdf)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = 'attachment; filename=research_paper.pdf'
        
        return response
    except Exception as e:
        # Log the error
        print(f"Error exporting PDF: {str(e)}")
        
        # Handle case when pdfkit or wkhtmltopdf is not installed
        if "wkhtmltopdf" in str(e).lower():
            return jsonify({
                'status': 'error',
                'message': 'PDF generation requires wkhtmltopdf. Please install it from the terminal using: sudo apt-get install wkhtmltopdf (Ubuntu/Debian) or brew install wkhtmltopdf (macOS)'
            }), 500
            
        # Return a general error response
        return jsonify({
            'status': 'error',
            'message': f'Failed to export PDF: {str(e)}'
        }), 500

# Alternative PDF export using pure HTML/CSS printing


@app.errorhandler(404)
def not_found(e):
    return render_template('not_found.html', page_title='Not Found'), 404

if __name__ == '__main__':
    app.run(debug=True)